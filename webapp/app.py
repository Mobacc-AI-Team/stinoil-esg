from __future__ import annotations

import os
import re
import mimetypes
from collections import defaultdict
from datetime import datetime, timezone
from dataclasses import dataclass
from functools import lru_cache
from html import escape
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

import psycopg
import requests
from flask import Flask, abort, g, redirect, render_template, request, url_for
from werkzeug.exceptions import HTTPException
from werkzeug.security import check_password_hash, generate_password_hash

from chatgpt_export_to_kb import format_yaml_list, slugify, yaml_escape

try:
    from docx import Document as DocxDocument
except Exception:  # noqa: BLE001
    DocxDocument = None

try:
    from pypdf import PdfReader
except Exception:  # noqa: BLE001
    PdfReader = None


WEBAPP_DIR = Path(__file__).resolve().parent
DEFAULT_KB_ROOT = WEBAPP_DIR.parent / "ChemieComplianceKennisbank"
RUNTIME_KB_ROOT = Path("/tmp") / "ChemieComplianceKennisbank"
URL_ALLOWED_SCHEMES = {"http", "https"}
FRONTMATTER_DELIMITER = "---"
SECTION_MAP = {
    "wetgeving": "01_Wetgeving",
    "vragen": "02_Vragen",
    "beoordelingen": "03_Beoordelingen",
    "antwoorden": "04_Antwoorden",
    "bronnen": "05_Bronnen",
    "locaties": "06_Locaties",
    "templates": "07_Templates",
    "index": "08_Index",
    "workflows": "09_Workflows",
}
ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
SOURCE_FOLDER_MAP = {
    ".pdf": "pdf",
    ".docx": "interne_documenten",
    ".txt": "interne_documenten",
    ".md": "interne_documenten",
}


@dataclass
class DocumentRecord:
    title: str
    path: Path
    rel_path: str
    section_key: str
    section_label: str
    metadata: dict[str, str]
    body: str

    @property
    def search_blob(self) -> str:
        meta_text = " ".join(f"{key} {value}" for key, value in self.metadata.items())
        return f"{self.title} {meta_text} {self.body}".lower()

    @property
    def tags(self) -> list[str]:
        raw = self.metadata.get("tags", "")
        if raw.startswith("[") and raw.endswith("]"):
            inner = raw[1:-1].strip()
            if not inner:
                return []
            return [clean_tag(part) for part in inner.split(",") if clean_tag(part)]
        tag = clean_tag(raw)
        return [tag] if tag else []

    @property
    def categories(self) -> list[str]:
        raw = self.metadata.get("categorieen", "")
        if raw.startswith("[") and raw.endswith("]"):
            inner = raw[1:-1].strip()
            if not inner:
                return []
            return [clean_tag(part) for part in inner.split(",") if clean_tag(part)]
        category = clean_tag(raw)
        return [category] if category else []


@dataclass
class DashboardStats:
    wetgeving: int
    vragen: int
    beoordelingen: int
    antwoorden: int
    locaties: int


@dataclass
class SearchFilters:
    query: str
    section: str
    tag: str
    category: str
    location: str
    status: str


@dataclass
class IntakePreview:
    afzender_type: str
    organisatie: str
    locatie: str
    titel: str
    bron: str
    tags: list[str]
    categories: list[str]
    suggested_regulations: list[str]
    intake_text: str
    attachment_name: str
    attachment_present: bool
    related_documents: list[dict[str, object]]


@dataclass
class AuthUser:
    user_id: int
    email: str
    role: str
    display_name: str


def create_app() -> Flask:
    app = Flask(
        __name__,
        template_folder=str(WEBAPP_DIR / "templates"),
        static_folder=str(WEBAPP_DIR / "static"),
        static_url_path="",
    )
    app.secret_key = os.environ.get("APP_SECRET_KEY", "dev-secret-change-me")
    kb_root = resolve_kb_root()
    app.config["KB_ROOT"] = kb_root
    app.config["DATABASE_URL"] = os.environ.get("POSTGRES_URL") or os.environ.get("DATABASE_URL", "").strip()

    if app.config["DATABASE_URL"]:
        ensure_database_schema(app.config["DATABASE_URL"])

    @app.before_request
    def load_user_into_context() -> None:
        g.current_user = get_current_user(app.config["DATABASE_URL"])

    @app.errorhandler(Exception)
    def handle_exception(exc: Exception):  # type: ignore[override]
        if isinstance(exc, HTTPException):
            return exc
        return (
            render_template(
                "error.html",
                error_type=type(exc).__name__,
                error_message=str(exc),
            ),
            500,
        )

    @app.errorhandler(401)
    def handle_unauthorized(_exc):  # type: ignore[override]
        return redirect(url_for("login", next=request.path))

    @app.errorhandler(403)
    def handle_forbidden(_exc):  # type: ignore[override]
        return render_template("forbidden.html"), 403

    @app.context_processor
    def inject_globals() -> dict[str, object]:
        return {
            "kb_root": kb_root,
            "section_options": list(SECTION_MAP.items()),
            "current_user": getattr(g, "current_user", None),
            "database_configured": bool(app.config["DATABASE_URL"]),
        }

    @app.route("/login", methods=["GET", "POST"])
    def login() -> str:
        if not app.config["DATABASE_URL"]:
            return render_template("login.html", error="Database is nog niet geconfigureerd. Voeg POSTGRES_URL of DATABASE_URL toe in Vercel.")
        if request.method == "POST":
            email = request.form.get("email", "").strip().lower()
            password = request.form.get("password", "")
            user = authenticate_user(app.config["DATABASE_URL"], email, password)
            if user is None:
                return render_template("login.html", error="Ongeldige inloggegevens.")
            session_set_user(user)
            next_url = request.args.get("next", "").strip()
            if next_url.startswith("/"):
                return redirect(next_url)
            return redirect(url_for("dashboard"))
        return render_template("login.html", error="")

    @app.route("/logout")
    def logout() -> str:
        session_clear_user()
        return redirect(url_for("login"))

    @app.route("/setup-admin")
    def setup_admin() -> str:
        if os.environ.get("ALLOW_ADMIN_BOOTSTRAP", "false").lower() != "true":
            abort(404)
        if not app.config["DATABASE_URL"]:
            return render_template(
                "bootstrap_done.html",
                email="compliance@stinoil.com",
                message="Database niet geconfigureerd. Voeg eerst POSTGRES_URL of DATABASE_URL toe in Vercel.",
            )
        ensure_admin_user(
            app.config["DATABASE_URL"],
            email="compliance@stinoil.com",
            password=os.environ.get("DEFAULT_ADMIN_PASSWORD", "ChangeMe123!"),
            role="admin",
            display_name="Compliance Beheer",
        )
        return render_template(
            "bootstrap_done.html",
            email="compliance@stinoil.com",
            message="Het beheeraccount is ingesteld. Je kunt nu inloggen.",
        )

    @app.route("/setup-check")
    def setup_check() -> str:
        database_url = app.config["DATABASE_URL"]
        database_connected = False
        users_table_present = False
        admin_present = False

        if database_url:
            conn = get_db_connection(database_url)
            if conn is not None:
                database_connected = True
                with conn:
                    with conn.cursor() as cur:
                        cur.execute(
                            "SELECT EXISTS (SELECT 1 FROM information_schema.tables WHERE table_name = 'app_users')"
                        )
                        users_table_present = bool(cur.fetchone()[0])
                        if users_table_present:
                            cur.execute(
                                "SELECT EXISTS (SELECT 1 FROM app_users WHERE email = %s)",
                                ("compliance@stinoil.com",),
                            )
                            admin_present = bool(cur.fetchone()[0])

        return render_template(
            "setup_check.html",
            database_configured=bool(database_url),
            database_connected=database_connected,
            users_table_present=users_table_present,
            admin_present=admin_present,
            bootstrap_enabled=os.environ.get("ALLOW_ADMIN_BOOTSTRAP", "false").lower() == "true",
        )

    @app.route("/")
    def dashboard() -> str:
        require_login(app.config["DATABASE_URL"])
        documents = load_documents(kb_root)
        stats = build_dashboard_stats(documents)
        recent_questions = [doc for doc in documents if doc.section_key == "vragen"][-8:]
        recent_answers = [doc for doc in documents if doc.section_key == "antwoorden"][-8:]
        return render_template(
            "dashboard.html",
            stats=stats,
            recent_questions=list(reversed(recent_questions)),
            recent_answers=list(reversed(recent_answers)),
            top_tags=top_tags(documents),
        )

    @app.route("/healthz")
    def healthz() -> tuple[dict[str, object], int]:
        return {
            "ok": True,
            "kb_root": str(kb_root),
            "kb_exists": kb_root.exists(),
            "documents": len(load_documents(kb_root)),
        }, 200

    @app.route("/wetgeving")
    def wetgeving() -> str:
        require_login(app.config["DATABASE_URL"])
        query = request.args.get("q", "").strip()
        documents = [doc for doc in load_documents(kb_root) if doc.section_key == "wetgeving"]
        if query:
            needle = query.lower()
            documents = [doc for doc in documents if needle in doc.search_blob]
        return render_template(
            "wetgeving.html",
            query=query,
            documents=documents,
            upload_status=request.args.get("status", "").strip(),
        )

    @app.route("/wetgeving/upload", methods=["GET", "POST"])
    def wetgeving_upload() -> str:
        require_role(app.config["DATABASE_URL"], {"admin", "editor"})
        if request.method == "POST":
            bron_url = request.form.get("bron_url", "").strip()
            if bron_url:
                try:
                    create_regulation_record_from_url(
                        kb_root=kb_root,
                        source_url=bron_url,
                        title=request.form.get("titel", "").strip(),
                        jurisdiction=request.form.get("jurisdictie", "").strip(),
                        subject=request.form.get("onderwerp", "").strip(),
                        tags=request.form.get("tags", "").strip(),
                        categories=request.form.get("categorieen", "").strip(),
                        source_label=request.form.get("bron", "").strip(),
                    )
                except ValueError:
                    return redirect(url_for("wetgeving", status="ongeldige_url"))
                except RuntimeError:
                    return redirect(url_for("wetgeving", status="url_import_mislukt"))

                load_documents.cache_clear()
                safe_build_indexes(kb_root)
                return redirect(url_for("wetgeving", status="toegevoegd"))

            upload = request.files.get("bestand")
            if upload is None or not upload.filename:
                return redirect(url_for("wetgeving", status="geen_bestand"))

            try:
                create_regulation_record_from_upload(
                    kb_root=kb_root,
                    upload_name=upload.filename,
                    binary_content=upload.read(),
                    title=request.form.get("titel", "").strip(),
                    jurisdiction=request.form.get("jurisdictie", "").strip(),
                    subject=request.form.get("onderwerp", "").strip(),
                    tags=request.form.get("tags", "").strip(),
                    categories=request.form.get("categorieen", "").strip(),
                    source_label=request.form.get("bron", "").strip(),
                )
            except ValueError:
                return redirect(url_for("wetgeving", status="ongeldig_bestand"))
            except RuntimeError:
                return redirect(url_for("wetgeving", status="extractie_mislukt"))

            load_documents.cache_clear()
            safe_build_indexes(kb_root)
            return redirect(url_for("wetgeving", status="toegevoegd"))

        return render_template("wetgeving_upload.html")

    @app.route("/casussen")
    def casussen() -> str:
        require_login(app.config["DATABASE_URL"])
        filters = SearchFilters(
            query=request.args.get("q", "").strip(),
            section=request.args.get("section", "vragen").strip(),
            tag=request.args.get("tag", "").strip(),
            category=request.args.get("category", "").strip(),
            location=request.args.get("location", "").strip(),
            status=request.args.get("status", "").strip(),
        )
        docs = filter_case_documents(load_documents(kb_root), filters)
        return render_template(
            "casussen.html",
            filters=filters,
            documents=docs,
            available_tags=sorted(unique_values(load_documents(kb_root), "tags")),
            available_categories=sorted(unique_values(load_documents(kb_root), "categorieen")),
            available_locations=sorted(unique_values(load_documents(kb_root), "locatie")),
            available_statuses=sorted(unique_values(load_documents(kb_root), "status")),
            create_status=request.args.get("status", "").strip(),
        )

    @app.route("/casussen/nieuw", methods=["GET", "POST"])
    def nieuwe_casus() -> str:
        require_role(app.config["DATABASE_URL"], {"admin", "editor"})
        ensure_case_directories(kb_root)
        if request.method == "POST":
            try:
                upload = request.files.get("bestand")
                action = request.form.get("actie", "preview").strip()
                form_payload = {
                    "afzender_type": request.form.get("afzender_type", "").strip(),
                    "organisatie": request.form.get("organisatie", "").strip(),
                    "locatie": request.form.get("locatie", "").strip(),
                    "titel": request.form.get("titel", "").strip(),
                    "vraag": request.form.get("vraag", "").strip(),
                    "tags": request.form.get("tags", "").strip(),
                    "categories": request.form.get("categorieen", "").strip(),
                    "bron": request.form.get("bron", "").strip(),
                    "upload_name": upload.filename if upload and upload.filename else request.form.get("upload_name", "").strip(),
                    "upload_content": upload.read() if upload and upload.filename else decode_hidden_bytes(request.form.get("upload_content", "")),
                }

                if action == "preview":
                    preview = build_case_preview(kb_root=kb_root, **form_payload)
                    template_values = {**form_payload, "upload_content": encode_hidden_bytes(form_payload["upload_content"])}
                    return render_template("casus_nieuw.html", preview=preview, form_values=template_values)

                form_payload["titel"] = request.form.get("preview_titel", form_payload["titel"]).strip()
                form_payload["afzender_type"] = request.form.get("preview_afzender_type", form_payload["afzender_type"]).strip()
                form_payload["organisatie"] = request.form.get("preview_organisatie", form_payload["organisatie"]).strip()
                form_payload["locatie"] = request.form.get("preview_locatie", form_payload["locatie"]).strip()
                form_payload["bron"] = request.form.get("preview_bron", form_payload["bron"]).strip()
                form_payload["vraag"] = request.form.get("preview_intake_text", form_payload["vraag"]).strip()
                form_payload["tags"] = request.form.get("preview_tags", form_payload["tags"]).strip()
                form_payload["categories"] = request.form.get("preview_categories", form_payload["categories"]).strip()

                create_case_from_form(kb_root=kb_root, **form_payload)
            except ValueError:
                return redirect(url_for("casussen", status="ongeldige_invoer"))

            load_documents.cache_clear()
            safe_build_indexes(kb_root)
            return redirect(url_for("casussen", status="toegevoegd", section="vragen"))

        return render_template("casus_nieuw.html", preview=None, form_values={})

    @app.route("/taxonomie")
    def taxonomie() -> str:
        require_login(app.config["DATABASE_URL"])
        documents = load_documents(kb_root)
        tree = build_taxonomy_tree(documents)
        return render_template("taxonomie.html", taxonomy=tree)

    @app.route("/zoek")
    def zoek() -> str:
        require_login(app.config["DATABASE_URL"])
        filters = SearchFilters(
            query=request.args.get("q", "").strip(),
            section=request.args.get("section", "").strip(),
            tag=request.args.get("tag", "").strip(),
            category=request.args.get("category", "").strip(),
            location=request.args.get("location", "").strip(),
            status=request.args.get("status", "").strip(),
        )
        docs = filter_documents(load_documents(kb_root), filters)
        return render_template("search.html", filters=filters, documents=docs)

    @app.route("/document/<path:rel_path>")
    def document_detail(rel_path: str) -> str:
        require_login(app.config["DATABASE_URL"])
        safe_rel = Path(rel_path)
        target = (kb_root / safe_rel).resolve()
        if kb_root not in [target, *target.parents] or not target.is_file():
            abort(404)
        record = document_from_file(kb_root, target)
        return render_template(
            "document_detail.html",
            document=record,
            rendered_body=markdown_to_html(record.body),
        )

    return app


@lru_cache(maxsize=1)
def load_documents(kb_root: Path) -> tuple[DocumentRecord, ...]:
    if not kb_root.exists():
        return tuple()

    records: list[DocumentRecord] = []
    for path in sorted(kb_root.glob("**/*.md")):
        records.append(document_from_file(kb_root, path))
    return tuple(records)


def resolve_kb_root() -> Path:
    configured = os.environ.get("KB_ROOT")
    if configured:
        target = Path(configured).expanduser().resolve()
        if target.exists() or ensure_writable_directory(target):
            return target

    default_root = DEFAULT_KB_ROOT.resolve()
    if is_writable_path(default_root):
        ensure_writable_directory(default_root)
        return default_root

    runtime_root = RUNTIME_KB_ROOT
    ensure_writable_directory(runtime_root)
    bootstrap_runtime_kb(default_root, runtime_root)
    return runtime_root


def is_writable_path(path: Path) -> bool:
    parent = path if path.exists() and path.is_dir() else path.parent
    try:
        parent.mkdir(parents=True, exist_ok=True)
        probe = parent / ".write_test"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
        return True
    except OSError:
        return False


def ensure_writable_directory(path: Path) -> bool:
    try:
        path.mkdir(parents=True, exist_ok=True)
        return is_writable_path(path)
    except OSError:
        return False


def bootstrap_runtime_kb(source_root: Path, runtime_root: Path) -> None:
    if any(runtime_root.iterdir()):
        return
    if not source_root.exists():
        return
    for path in source_root.glob("**/*"):
        relative = path.relative_to(source_root)
        target = runtime_root / relative
        if path.is_dir():
            target.mkdir(parents=True, exist_ok=True)
        elif path.is_file():
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(path.read_bytes())


def document_from_file(kb_root: Path, path: Path) -> DocumentRecord:
    metadata, body = split_frontmatter(path.read_text(encoding="utf-8"))
    rel_path = str(path.relative_to(kb_root)).replace("\\", "/")
    section_key = infer_section_key(rel_path)
    return DocumentRecord(
        title=metadata.get("titel") or path.stem.replace("_", " ").title(),
        path=path,
        rel_path=rel_path,
        section_key=section_key,
        section_label=section_label(section_key),
        metadata=metadata,
        body=body,
    )


def split_frontmatter(text: str) -> tuple[dict[str, str], str]:
    if not text.startswith(FRONTMATTER_DELIMITER):
        return {}, text

    lines = text.splitlines()
    metadata: dict[str, str] = {}
    end_index = 0
    for index, line in enumerate(lines[1:], start=1):
        if line.strip() == FRONTMATTER_DELIMITER:
            end_index = index
            break
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        metadata[key.strip()] = value.strip()

    body = "\n".join(lines[end_index + 1 :]).strip()
    return metadata, body


def infer_section_key(rel_path: str) -> str:
    top = rel_path.split("/", 1)[0]
    for key, folder in SECTION_MAP.items():
        if folder == top:
            return key
    return "overig"


def section_label(section_key: str) -> str:
    return {
        "wetgeving": "Wetgeving",
        "vragen": "Vragen",
        "beoordelingen": "Beoordelingen",
        "antwoorden": "Antwoorden",
        "bronnen": "Bronnen",
        "locaties": "Locaties",
        "templates": "Templates",
        "index": "Index",
        "workflows": "Workflows",
        "overig": "Overig",
    }.get(section_key, section_key.title())


def build_dashboard_stats(documents: Iterable[DocumentRecord]) -> DashboardStats:
    docs = list(documents)
    return DashboardStats(
        wetgeving=count_section(docs, "wetgeving"),
        vragen=count_section(docs, "vragen"),
        beoordelingen=count_section(docs, "beoordelingen"),
        antwoorden=count_section(docs, "antwoorden"),
        locaties=count_section(docs, "locaties"),
    )


def count_section(documents: Iterable[DocumentRecord], section_key: str) -> int:
    return sum(1 for doc in documents if doc.section_key == section_key)


def top_tags(documents: Iterable[DocumentRecord]) -> list[tuple[str, int]]:
    tag_counts: dict[str, int] = {}
    for doc in documents:
        for tag in doc.tags:
            tag_counts[tag] = tag_counts.get(tag, 0) + 1
    return sorted(tag_counts.items(), key=lambda item: (-item[1], item[0]))[:12]


def build_taxonomy_tree(documents: Iterable[DocumentRecord]) -> list[dict[str, object]]:
    category_map: dict[str, list[DocumentRecord]] = defaultdict(list)
    for doc in documents:
        for category in doc.categories:
            category_map[category].append(doc)

    taxonomy: list[dict[str, object]] = []
    for category in sorted(category_map):
        docs = sorted(category_map[category], key=lambda item: (item.section_label, item.title.casefold()))
        taxonomy.append(
            {
                "name": category,
                "count": len(docs),
                "documents": docs,
            }
        )
    return taxonomy


def unique_values(documents: Iterable[DocumentRecord], field: str) -> set[str]:
    results: set[str] = set()
    for doc in documents:
        if field == "tags":
            results.update(doc.tags)
            continue
        if field == "categorieen":
            results.update(doc.categories)
            continue
        value = doc.metadata.get(field, "").strip().strip('"')
        if value:
            results.add(value)
    return results


def filter_documents(documents: Iterable[DocumentRecord], filters: SearchFilters) -> list[DocumentRecord]:
    results = list(documents)
    if filters.section:
        results = [doc for doc in results if doc.section_key == filters.section]
    if filters.query:
        needle = filters.query.lower()
        results = [doc for doc in results if needle in doc.search_blob]
    if filters.tag:
        results = [doc for doc in results if filters.tag in doc.tags]
    if filters.category:
        results = [doc for doc in results if filters.category in doc.categories]
    if filters.location:
        results = [doc for doc in results if doc.metadata.get("locatie", "").strip('"') == filters.location]
    if filters.status:
        results = [doc for doc in results if doc.metadata.get("status", "").strip('"') == filters.status]
    return results


def filter_case_documents(documents: Iterable[DocumentRecord], filters: SearchFilters) -> list[DocumentRecord]:
    case_sections = {"vragen", "beoordelingen", "antwoorden"}
    docs = [doc for doc in documents if doc.section_key in case_sections]
    if filters.section and filters.section in case_sections:
        docs = [doc for doc in docs if doc.section_key == filters.section]
    if filters.query:
        needle = filters.query.lower()
        docs = [doc for doc in docs if needle in doc.search_blob]
    if filters.tag:
        docs = [doc for doc in docs if filters.tag in doc.tags]
    if filters.category:
        docs = [doc for doc in docs if filters.category in doc.categories]
    if filters.location:
        docs = [doc for doc in docs if doc.metadata.get("locatie", "").strip('"') == filters.location]
    if filters.status:
        docs = [doc for doc in docs if doc.metadata.get("status", "").strip('"') == filters.status]
    return docs


def clean_tag(value: str) -> str:
    return value.strip().strip('"').strip("'")


def create_regulation_record_from_upload(
    kb_root: Path,
    upload_name: str,
    binary_content: bytes,
    title: str,
    jurisdiction: str,
    subject: str,
    tags: str,
    categories: str,
    source_label: str,
) -> Path:
    suffix = Path(upload_name).suffix.lower()
    if suffix not in ALLOWED_UPLOAD_EXTENSIONS:
        raise ValueError("Unsupported file type")

    safe_title = title or Path(upload_name).stem.replace("_", " ").strip() or "Wetgeving upload"
    normalized_jurisdiction = normalize_jurisdiction(jurisdiction)
    target_slug = slugify(safe_title)
    timestamp = datetime.utcnow().strftime("%Y-%m-%d")

    source_dir = kb_root / "05_Bronnen" / SOURCE_FOLDER_MAP[suffix]
    source_dir.mkdir(parents=True, exist_ok=True)
    source_path = unique_path(source_dir / f"{target_slug}{suffix}")
    source_path.write_bytes(binary_content)

    extracted_text = extract_text_from_file(source_path)
    regulation_dir = kb_root / "01_Wetgeving" / normalized_jurisdiction
    regulation_dir.mkdir(parents=True, exist_ok=True)
    regulation_path = unique_path(regulation_dir / f"{target_slug}.md")

    tag_list = parse_tags(tags)
    if not tag_list:
        tag_list = derive_tags_from_text(f"{safe_title} {subject} {extracted_text[:4000]}")

    category_list = parse_tags(categories)
    if not category_list:
        category_list = derive_categories_from_text(f"{safe_title} {subject} {extracted_text[:6000]}")

    summary = summarize_extracted_text(extracted_text, subject or safe_title)
    source_rel = str(source_path.relative_to(kb_root)).replace("\\", "/")
    regulation_path.write_text(
        render_uploaded_regulation_markdown(
            title=safe_title,
            jurisdiction=normalized_jurisdiction,
            subject=subject or infer_subject(extracted_text),
            source_label=source_label or upload_name,
            source_rel=source_rel,
            summary=summary,
            tags=tag_list,
            categories=category_list,
            extracted_text=extracted_text,
            timestamp=timestamp,
        ),
        encoding="utf-8",
    )
    return regulation_path


def strip_html_to_text(html: str) -> str:
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def create_regulation_record_from_url(
    kb_root: Path,
    source_url: str,
    title: str,
    jurisdiction: str,
    subject: str,
    tags: str,
    categories: str,
    source_label: str,
) -> Path:
    parsed = urlparse(source_url)
    if parsed.scheme not in URL_ALLOWED_SCHEMES:
        raise ValueError("Alleen http/https URLs zijn toegestaan")

    response = requests.get(source_url, timeout=30)
    response.raise_for_status()
    content_type = response.headers.get("content-type", "text/html").split(";", 1)[0].strip()
    suffix = infer_suffix_from_response(content_type, parsed.path)

    if "html" in content_type:
        binary_content = strip_html_to_text(response.text).encode("utf-8")
    else:
        binary_content = response.content

    upload_name = f"{slugify(title or Path(parsed.path).stem or 'wetgeving_url')}{suffix}"
    return create_regulation_record_from_upload(
        kb_root=kb_root,
        upload_name=upload_name,
        binary_content=binary_content,
        title=title,
        jurisdiction=jurisdiction,
        subject=subject,
        tags=tags,
        categories=categories,
        source_label=source_label or source_url,
    )


def create_case_from_form(
    kb_root: Path,
    afzender_type: str,
    organisatie: str,
    locatie: str,
    titel: str,
    vraag: str,
    tags: str,
    categories: str,
    bron: str,
    upload_name: str,
    upload_content: bytes,
) -> tuple[Path, Path, Path]:
    ensure_case_directories(kb_root)
    preview = build_case_preview(
        kb_root=kb_root,
        afzender_type=afzender_type,
        organisatie=organisatie,
        locatie=locatie,
        titel=titel,
        vraag=vraag,
        tags=tags,
        categories=categories,
        bron=bron,
        upload_name=upload_name,
        upload_content=upload_content,
    )

    now = datetime.now(timezone.utc)
    year = now.strftime("%Y")
    date_str = now.strftime("%Y-%m-%d")
    slug = slugify(preview.titel)
    if organisatie:
        slug = f"{slugify(organisatie)}_{slug}"

    vraag_dir = kb_root / "02_Vragen" / year
    beoordeling_dir = kb_root / "03_Beoordelingen" / year
    antwoord_dir = kb_root / "04_Antwoorden" / year
    vraag_dir.mkdir(parents=True, exist_ok=True)
    beoordeling_dir.mkdir(parents=True, exist_ok=True)
    antwoord_dir.mkdir(parents=True, exist_ok=True)

    vraag_path, beoordeling_path, antwoord_path = unique_case_paths(
        vraag_dir, beoordeling_dir, antwoord_dir, date_str, slug
    )

    tag_list = preview.tags
    category_list = preview.categories
    summary = summarize_extracted_text(preview.intake_text, preview.titel, max_len=220)
    suggested_regulations = preview.suggested_regulations
    vraag_rel = relative_path(kb_root, vraag_path)
    beoordeling_rel = relative_path(kb_root, beoordeling_path)
    antwoord_rel = relative_path(kb_root, antwoord_path)
    attachment_rel = ""
    if preview.attachment_present and upload_name and upload_content:
        attachment_path = save_case_attachment(kb_root, upload_name, upload_content, organisatie or preview.titel)
        attachment_rel = relative_path(kb_root, attachment_path)

    vraag_path.write_text(
        render_case_question_markdown(
            titel=preview.titel,
            datum=date_str,
            afzender_type=preview.afzender_type,
            organisatie=organisatie or "onbekend",
            locatie=locatie or "centrale_beoordeling",
            vraag=preview.intake_text,
            tags=tag_list,
            categories=category_list,
            bron=preview.bron,
            attachment_rel=attachment_rel,
            beoordeling_rel=beoordeling_rel,
            antwoord_rel=antwoord_rel,
            summary=summary,
        ),
        encoding="utf-8",
    )

    beoordeling_path.write_text(
        render_case_assessment_markdown(
            titel=preview.titel,
            datum=date_str,
            vraag_rel=vraag_rel,
            antwoord_rel=antwoord_rel,
            tags=tag_list,
            categories=category_list,
            summary=summary,
            vraag=preview.intake_text,
            suggested_regulations=suggested_regulations,
        ),
        encoding="utf-8",
    )

    antwoord_path.write_text(
        render_case_answer_markdown(
            titel=preview.titel,
            datum=date_str,
            doelgroep=preview.afzender_type,
            vraag_rel=vraag_rel,
            beoordeling_rel=beoordeling_rel,
            tags=tag_list,
            categories=category_list,
            summary=summary,
            suggested_regulations=suggested_regulations,
        ),
        encoding="utf-8",
    )

    return vraag_path, beoordeling_path, antwoord_path


def safe_build_indexes(kb_root: Path) -> None:
    try:
        from chatgpt_export_to_kb import build_indexes
    except Exception:
        return
    try:
        build_indexes(kb_root)
    except Exception:
        return


def get_db_connection(database_url: str):
    if not database_url:
        return None
    try:
        return psycopg.connect(database_url)
    except Exception:
        return None


def ensure_database_schema(database_url: str) -> None:
    conn = get_db_connection(database_url)
    if conn is None:
        return
    with conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS app_users (
                    id SERIAL PRIMARY KEY,
                    email TEXT UNIQUE NOT NULL,
                    password_hash TEXT NOT NULL,
                    role TEXT NOT NULL DEFAULT 'viewer',
                    display_name TEXT NOT NULL,
                    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                )
                """
            )
            cur.execute("SELECT COUNT(*) FROM app_users")
            if cur.fetchone()[0] == 0:
                seed_default_users(cur)


def seed_default_users(cur: Any) -> None:
    default_password = os.environ.get("DEFAULT_ADMIN_PASSWORD", "ChangeMe123!")
    users = [
        ("compliance@stinoil.com", generate_password_hash(default_password), "admin", "Compliance Beheer"),
        ("editor@example.com", generate_password_hash(default_password), "editor", "Editor"),
        ("viewer@example.com", generate_password_hash(default_password), "viewer", "Viewer"),
    ]
    cur.executemany(
        "INSERT INTO app_users (email, password_hash, role, display_name) VALUES (%s, %s, %s, %s)",
        users,
    )


def ensure_admin_user(database_url: str, email: str, password: str, role: str, display_name: str) -> None:
    conn = get_db_connection(database_url)
    if conn is None:
        raise RuntimeError("Database niet beschikbaar")
    with conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id FROM app_users WHERE email = %s", (email,))
            row = cur.fetchone()
            password_hash = generate_password_hash(password)
            if row is None:
                cur.execute(
                    "INSERT INTO app_users (email, password_hash, role, display_name) VALUES (%s, %s, %s, %s)",
                    (email, password_hash, role, display_name),
                )
            else:
                cur.execute(
                    "UPDATE app_users SET password_hash = %s, role = %s, display_name = %s WHERE email = %s",
                    (password_hash, role, display_name, email),
                )


def authenticate_user(database_url: str, email: str, password: str) -> AuthUser | None:
    conn = get_db_connection(database_url)
    if conn is None:
        return None
    with conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, email, password_hash, role, display_name FROM app_users WHERE email = %s",
                (email,),
            )
            row = cur.fetchone()
            if row is None or not check_password_hash(row[2], password):
                return None
            return AuthUser(user_id=row[0], email=row[1], role=row[3], display_name=row[4])


def get_current_user(database_url: str) -> AuthUser | None:
    from flask import session

    user_id = session.get("user_id")
    if not user_id or not database_url:
        return None
    conn = get_db_connection(database_url)
    if conn is None:
        return None
    with conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, email, role, display_name FROM app_users WHERE id = %s",
                (user_id,),
            )
            row = cur.fetchone()
            if row is None:
                return None
            return AuthUser(user_id=row[0], email=row[1], role=row[2], display_name=row[3])


def require_login(database_url: str) -> None:
    if not database_url:
        return
    if get_current_user(database_url) is None:
        abort(401)


def require_role(database_url: str, allowed_roles: set[str]) -> None:
    if not database_url:
        return
    user = get_current_user(database_url)
    if user is None:
        abort(401)
    if user.role not in allowed_roles:
        abort(403)


def session_set_user(user: AuthUser) -> None:
    from flask import session

    session["user_id"] = user.user_id


def session_clear_user() -> None:
    from flask import session

    session.pop("user_id", None)


def build_case_preview(
    kb_root: Path,
    afzender_type: str,
    organisatie: str,
    locatie: str,
    titel: str,
    vraag: str,
    tags: str,
    categories: str,
    bron: str,
    upload_name: str,
    upload_content: bytes,
) -> IntakePreview:
    normalized_type = slugify(afzender_type) if afzender_type else "extern"
    if normalized_type not in {"klant", "leverancier", "extern"}:
        raise ValueError("Ongeldig afzendertype")

    extracted_upload_text = ""
    if upload_name and upload_content:
        temp_path = write_temp_upload(upload_name, upload_content)
        try:
            extracted_upload_text = extract_text_from_file(temp_path)
        finally:
            temp_path.unlink(missing_ok=True)

    intake_text = question_intake_text(vraag, extracted_upload_text)
    resolved_title = titel or infer_case_title(intake_text, organisatie, normalized_type)
    if not resolved_title or not intake_text:
        raise ValueError("Titel en vraag zijn verplicht")

    tag_list = parse_tags(tags)
    if not tag_list:
        tag_list = derive_tags_from_text(f"{resolved_title} {intake_text}")

    category_list = parse_tags(categories)
    if not category_list:
        category_list = derive_categories_from_text(f"{resolved_title} {intake_text}")

    return IntakePreview(
        afzender_type=normalized_type,
        organisatie=organisatie or "onbekend",
        locatie=locatie or "centrale_beoordeling",
        titel=resolved_title,
        bron=bron or "webformulier",
        tags=tag_list,
        categories=category_list,
        suggested_regulations=suggest_relevant_regulations(tag_list, category_list, intake_text),
        intake_text=intake_text,
        attachment_name=upload_name,
        attachment_present=bool(upload_name and upload_content),
        related_documents=find_related_documents(kb_root, resolved_title, intake_text, tag_list, category_list),
    )


def find_related_documents(
    kb_root: Path,
    title: str,
    intake_text: str,
    tags: list[str],
    categories: list[str],
    limit: int = 6,
) -> list[dict[str, object]]:
    documents = load_documents(kb_root)
    tokens = tokenize_similarity_text(f"{title} {intake_text} {' '.join(tags)} {' '.join(categories)}")
    results: list[dict[str, object]] = []
    for doc in documents:
        if doc.section_key not in {"vragen", "beoordelingen", "antwoorden", "wetgeving"}:
            continue
        doc_tokens = tokenize_similarity_text(
            f"{doc.title} {doc.body} {' '.join(doc.tags)} {' '.join(doc.categories)}"
        )
        score = similarity_score(tokens, doc_tokens, tags, doc.tags, categories, doc.categories)
        if score <= 0:
            continue
        results.append(
            {
                "title": doc.title,
                "rel_path": doc.rel_path,
                "section_label": doc.section_label,
                "score": score,
                "tags": doc.tags,
                "categories": doc.categories,
                "summary": doc.metadata.get("samenvatting", ""),
            }
        )

    results.sort(key=lambda item: (-float(item["score"]), str(item["title"]).casefold()))
    return results[:limit]


def tokenize_similarity_text(text: str) -> set[str]:
    words = re.findall(r"[a-zA-Z0-9_]{3,}", text.lower())
    stopwords = {
        "de", "het", "een", "voor", "van", "met", "over", "aan", "bij", "die", "dit",
        "wordt", "zijn", "naar", "nog", "ook", "als", "dan", "door", "from", "that",
    }
    return {word for word in words if word not in stopwords}


def similarity_score(
    input_tokens: set[str],
    doc_tokens: set[str],
    input_tags: list[str],
    doc_tags: list[str],
    input_categories: list[str],
    doc_categories: list[str],
) -> float:
    overlap = len(input_tokens & doc_tokens)
    tag_overlap = len(set(input_tags) & set(doc_tags))
    category_overlap = len(set(input_categories) & set(doc_categories))
    return overlap + (tag_overlap * 4) + (category_overlap * 5)


def write_temp_upload(upload_name: str, upload_content: bytes) -> Path:
    suffix = Path(upload_name).suffix.lower()
    if suffix not in ALLOWED_UPLOAD_EXTENSIONS:
        raise ValueError("Unsupported file type")
    temp_dir = WEBAPP_DIR.parent / ".tmp_uploads"
    temp_dir.mkdir(parents=True, exist_ok=True)
    path = unique_path(temp_dir / f"preview{suffix}")
    path.write_bytes(upload_content)
    return path


def infer_suffix_from_response(content_type: str, path: str) -> str:
    if "pdf" in content_type:
        return ".pdf"
    if "wordprocessingml" in content_type:
        return ".docx"
    if "html" in content_type:
        return ".txt"
    suffix = Path(path).suffix.lower()
    if suffix in ALLOWED_UPLOAD_EXTENSIONS:
        return suffix
    return ".txt"


def encode_hidden_bytes(value: bytes) -> str:
    return value.hex() if value else ""


def decode_hidden_bytes(value: str) -> bytes:
    return bytes.fromhex(value) if value else b""


def ensure_case_directories(kb_root: Path) -> None:
    for rel in ["02_Vragen", "03_Beoordelingen", "04_Antwoorden", "05_Bronnen/emails", "05_Bronnen/pdf", "05_Bronnen/interne_documenten"]:
        (kb_root / rel).mkdir(parents=True, exist_ok=True)


def save_case_attachment(kb_root: Path, upload_name: str, upload_content: bytes, seed: str) -> Path:
    suffix = Path(upload_name).suffix.lower()
    if suffix not in ALLOWED_UPLOAD_EXTENSIONS:
        raise ValueError("Unsupported file type")
    folder = SOURCE_FOLDER_MAP.get(suffix, "interne_documenten")
    base_slug = slugify(seed or Path(upload_name).stem or "bijlage")
    target_dir = kb_root / "05_Bronnen" / folder
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = unique_path(target_dir / f"{base_slug}{suffix}")
    target_path.write_bytes(upload_content)
    return target_path


def question_intake_text(manual_text: str, extracted_upload_text: str) -> str:
    parts = [part.strip() for part in [manual_text, extracted_upload_text] if part and part.strip()]
    return "\n\n".join(parts)


def infer_case_title(intake_text: str, organisatie: str, afzender_type: str) -> str:
    base = summarize_extracted_text(intake_text, "Nieuwe vraag", max_len=90)
    prefix = organisatie or afzender_type or "casus"
    title = f"{prefix} - {base}".strip(" -")
    return title[:120].strip()


def suggest_relevant_regulations(tags: list[str], categories: list[str], intake_text: str) -> list[str]:
    text = f"{' '.join(tags)} {' '.join(categories)} {intake_text}".lower()
    suggestions: list[str] = []
    rules = {
        "REACH-verordening": ["reach", "svhc", "sds"],
        "CLP-verordening": ["clp", "etikettering", "classificatie"],
        "ADR": ["adr", "transport", "un-nummer"],
        "PGS 15": ["pgs", "opslag"],
        "Brzo / Seveso": ["brzo", "seveso", "zware ongevallen"],
        "Arbeidsomstandighedenbesluit": ["arbeidsomstandighedenbesluit", "arbobesluit"],
        "Arbeidsomstandighedenregeling": ["arbeidsomstandighedenregeling", "arboregeling"],
        "ATEX": ["atex", "explosieve atmosfeer"],
        "QRA": ["qra", "risicoanalyse"],
        "Vergunningen / Omgevingswet": ["vergunning", "omgevingsvergunning", "bal"],
    }
    for name, markers in rules.items():
        if any(marker in text for marker in markers):
            suggestions.append(name)
    return suggestions or ["Nog handmatig te koppelen aan relevante wetgeving"]


def unique_case_paths(
    vraag_dir: Path,
    beoordeling_dir: Path,
    antwoord_dir: Path,
    date_str: str,
    slug: str,
) -> tuple[Path, Path, Path]:
    counter = 1
    while True:
        suffix = "" if counter == 1 else f"_{counter}"
        base = f"{date_str}_{slug}{suffix}"
        vraag_path = vraag_dir / f"{base}_vraag.md"
        beoordeling_path = beoordeling_dir / f"{base}_beoordeling.md"
        antwoord_path = antwoord_dir / f"{base}_antwoord.md"
        if not any(path.exists() for path in (vraag_path, beoordeling_path, antwoord_path)):
            return vraag_path, beoordeling_path, antwoord_path
        counter += 1


def render_case_question_markdown(
    titel: str,
    datum: str,
    afzender_type: str,
    organisatie: str,
    locatie: str,
    vraag: str,
    tags: list[str],
    categories: list[str],
    bron: str,
    attachment_rel: str,
    beoordeling_rel: str,
    antwoord_rel: str,
    summary: str,
) -> str:
    attachment_block = f"- Bijlage: `{attachment_rel}`" if attachment_rel else "- Geen bijlage opgeslagen"
    return f"""---
titel: {yaml_escape(titel)}
datum: {datum}
locatie: {yaml_escape(locatie)}
afzender_type: {yaml_escape(afzender_type)}
organisatie: {yaml_escape(organisatie)}
status: nieuw
bron: {yaml_escape(bron)}
tags: {format_yaml_list(tags)}
categorieen: {format_yaml_list(categories)}
beoordeling: {yaml_escape(beoordeling_rel)}
antwoord: {yaml_escape(antwoord_rel)}
samenvatting: {yaml_escape(summary)}
---

# Vraag

## Onderwerp
{titel}

## Herkomst
- Type afzender: {afzender_type}
- Organisatie: {organisatie}
- Locatie: {locatie}
- Invoerbron: {bron}
- Bijlagen: zie bronnen

## Vraagstelling
{vraag}

## Bijlagen / bronverwijzingen
{attachment_block}
""".strip() + "\n"


def render_case_assessment_markdown(
    titel: str,
    datum: str,
    vraag_rel: str,
    antwoord_rel: str,
    tags: list[str],
    categories: list[str],
    summary: str,
    vraag: str,
    suggested_regulations: list[str],
) -> str:
    category_block = "\n".join(f"- {item}" for item in categories) if categories else "- Nog te bepalen"
    regulation_block = "\n".join(f"- {item}" for item in suggested_regulations) if suggested_regulations else "- Nog aan te vullen"
    return f"""---
titel: {yaml_escape(titel)}
datum: {datum}
status: concept
tags: {format_yaml_list(tags)}
categorieen: {format_yaml_list(categories)}
vraagbestand: {yaml_escape(vraag_rel)}
antwoordbestand: {yaml_escape(antwoord_rel)}
eigenaar: compliance
samenvatting: {yaml_escape(summary)}
---

# Beoordeling

## Casus
{titel}

## Samenvatting
{summary}

## Categorieen
{category_block}

## Vraag
{vraag}

## Relevante wetgeving / normen / kaders
{regulation_block}

## Analyse
- Nog aan te vullen

## Uniform standpunt
- Nog vast te stellen
""".strip() + "\n"


def render_case_answer_markdown(
    titel: str,
    datum: str,
    doelgroep: str,
    vraag_rel: str,
    beoordeling_rel: str,
    tags: list[str],
    categories: list[str],
    summary: str,
    suggested_regulations: list[str],
) -> str:
    regulation_block = "\n".join(f"- {item}" for item in suggested_regulations) if suggested_regulations else "- Nog aan te vullen"
    return f"""---
titel: {yaml_escape(titel)}
datum: {datum}
doelgroep: {yaml_escape(doelgroep)}
status: concept
tags: {format_yaml_list(tags)}
categorieen: {format_yaml_list(categories)}
vraagbestand: {yaml_escape(vraag_rel)}
beoordelingsbestand: {yaml_escape(beoordeling_rel)}
samenvatting: {yaml_escape(summary)}
---

# Antwoord

## Onderwerp
{titel}

## Conceptantwoord
Nog op te stellen.

## Onderbouwing
- Toets aan beoordeling: `{beoordeling_rel}`
- Controle op locatiespecifieke afwijkingen: nog uitvoeren

## Eerste relevante kaders
{regulation_block}
""".strip() + "\n"


def normalize_jurisdiction(value: str) -> str:
    cleaned = slugify(value) if value else "nationaal"
    if cleaned in {"eu", "nationaal", "lokaal", "normen_en_richtlijnen"}:
        return cleaned
    return "nationaal"


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    counter = 2
    while True:
        candidate = path.with_name(f"{path.stem}_{counter}{path.suffix}")
        if not candidate.exists():
            return candidate
        counter += 1


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        if PdfReader is None:
            raise RuntimeError("pypdf ontbreekt")
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()
    if suffix == ".docx":
        if DocxDocument is None:
            raise RuntimeError("python-docx ontbreekt")
        document = DocxDocument(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
    raise ValueError("Unsupported file type")


def parse_tags(raw: str) -> list[str]:
    return [clean_tag(part) for part in raw.split(",") if clean_tag(part)]


def derive_tags_from_text(text: str) -> list[str]:
    lowered = text.lower()
    vocabulary = [
        "reach",
        "clp",
        "adr",
        "pgs15",
        "brzo",
        "bal",
        "etikettering",
        "sds",
        "veiligheid",
        "gevaarlijke stoffen",
        "emissie",
        "opslag",
        "transport",
    ]
    found = [slugify(tag) if " " in tag else tag for tag in vocabulary if tag in lowered]
    return found or ["wetgeving_upload"]


def derive_categories_from_text(text: str) -> list[str]:
    lowered = text.lower()
    category_rules = {
        "brzo": ["brzo", "seveso", "zware ongevallen"],
        "atex": ["atex", "explosieve atmosfeer", "explosiegevaar"],
        "arbeidsomstandighedenbesluit": ["arbeidsomstandighedenbesluit", "arbobesluit"],
        "arbeidsomstandighedenregeling": ["arbeidsomstandighedenregeling", "arboregeling"],
        "qra": ["qra", "quantitative risk assessment", "risicoanalyse"],
        "pgs_richtlijnen": ["pgs", "publicatiereeks gevaarlijke stoffen"],
        "vergunningen": ["vergunning", "omgevingsvergunning", "maatwerkvoorschrift"],
        "reach": ["reach", "svhc", "restrictie", "autorisatie"],
        "clp": ["clp", "etikettering", "classificatie", "gevarenpictogram"],
        "adr": ["adr", "transport", "un-nummer", "vervoersdocument"],
        "sds": ["sds", "veiligheidsinformatieblad"],
        "opslag_gevaarlijke_stoffen": ["opslag", "magazijn", "incompatibele stoffen"],
        "emissies": ["emissie", "lucht", "waterlozing", "vos"],
    }
    found: list[str] = []
    for category, markers in category_rules.items():
        if any(marker in lowered for marker in markers):
            found.append(category)
    return found or ["algemene_wetgeving"]


def summarize_extracted_text(text: str, fallback: str, max_len: int = 280) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return fallback
    return cleaned[: max_len - 1].rstrip() + "…" if len(cleaned) > max_len else cleaned


def infer_subject(text: str) -> str:
    lowered = text.lower()
    if "etiket" in lowered or "label" in lowered:
        return "etikettering"
    if "transport" in lowered or "adr" in lowered:
        return "transport"
    if "opslag" in lowered or "pgs" in lowered:
        return "opslag gevaarlijke stoffen"
    if "veiligheidsinformatieblad" in lowered or "sds" in lowered:
        return "veiligheidsinformatieblad"
    return "nog_te_bepalen"


def render_uploaded_regulation_markdown(
    title: str,
    jurisdiction: str,
    subject: str,
    source_label: str,
    source_rel: str,
    summary: str,
    tags: list[str],
    categories: list[str],
    extracted_text: str,
    timestamp: str,
) -> str:
    searchable_text = extracted_text.strip() or "Geen extraheerbare tekst gevonden."
    categories_block = "\n".join(f"- {category}" for category in categories) if categories else "- Nog te bepalen"
    return f"""---
titel: {yaml_escape(title)}
datum: {timestamp}
jurisdictie: {yaml_escape(jurisdiction)}
onderwerp: {yaml_escape(subject)}
bron: {yaml_escape(source_label)}
bronbestand: {yaml_escape(source_rel)}
status: actief
tags: {format_yaml_list(tags)}
categorieen: {format_yaml_list(categories)}
samenvatting: {yaml_escape(summary)}
---

# Wetgeving

## Onderwerp
{title}

## Reikwijdte
Nog te valideren op basis van de brontekst.

## Relevantie voor het bedrijf
Nog aan te vullen na inhoudelijke beoordeling.

## Categorieen
{categories_block}

## Kernverplichtingen
- Nog aan te vullen

## Toetsingspunten
- Nog aan te vullen

## Bronnen
- Upload: `{source_rel}`

## Geextraheerde brontekst
{searchable_text}
""".strip() + "\n"


def markdown_to_html(markdown_text: str) -> str:
    blocks = markdown_text.split("\n\n")
    html_blocks: list[str] = []
    for block in blocks:
        stripped = block.strip()
        if not stripped:
            continue
        lines = stripped.splitlines()
        if all(line.lstrip().startswith("- ") for line in lines):
            items = "".join(f"<li>{inline_markdown(line.lstrip()[2:])}</li>" for line in lines)
            html_blocks.append(f"<ul>{items}</ul>")
            continue
        if stripped.startswith("### "):
            html_blocks.append(f"<h3>{inline_markdown(stripped[4:])}</h3>")
            continue
        if stripped.startswith("## "):
            html_blocks.append(f"<h2>{inline_markdown(stripped[3:])}</h2>")
            continue
        if stripped.startswith("# "):
            html_blocks.append(f"<h1>{inline_markdown(stripped[2:])}</h1>")
            continue
        html_blocks.append("".join(f"<p>{inline_markdown(line)}</p>" for line in lines))
    return "\n".join(html_blocks)


def inline_markdown(text: str) -> str:
    text = escape(text)
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="/document/\2">\1</a>', text)
    return text


app = create_app()


if __name__ == "__main__":
    app.run(debug=True)
